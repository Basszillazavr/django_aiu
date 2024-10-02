import openai
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document

openai.api_key = 'ВАШ_API_КЛЮЧ'

def generate_doc(request):
    if request.method == 'POST':
        doc_type = request.POST.get('doc_type')
        topic = request.POST.get('topic')

        # Генерация текста с GPT
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=f"Напишите {doc_type} на тему {topic}",
            max_tokens=1500
        )
        generated_text = response['choices'][0]['text']

        # Создание документа
        document = Document()
        document.add_heading(f'{doc_type.capitalize()} на тему "{topic}"', 0)
        document.add_paragraph(generated_text)

        # Сохранение файла
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{doc_type}_{topic}.docx"'
        document.save(response)

        return response

    return render(request, 'index.html')
